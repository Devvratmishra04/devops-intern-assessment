import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def create_document():
    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("SAST Delivery & Rollout Automated Reporting (DEV-460)")
    run_title.font.name = "Calibri"
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(16, 44, 87)

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(14)
    run_sub = sub_p.add_run("Automated Scheduled Status Delivery, Rollout Roadmaps & Risk Mitigations")
    run_sub.font.name = "Calibri"
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(80, 80, 80)

    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Ticket ID:", "DEV-460 (https://glynac.atlassian.net/browse/DEV-460)"),
        ("Delivery Status:", "Done / Resolved on 2026-08-31"),
        ("Owner / Team:", "@DevOps"),
        ("Schedule:", "Every Monday at 09:00 AM UTC (cron: 0 9 * * 1)")
    ]
    for i, (k, v) in enumerate(meta_data):
        cell_k = meta_table.cell(i, 0)
        cell_v = meta_table.cell(i, 1)
        cell_k.text = k
        cell_v.text = v
        cell_k.paragraphs[0].runs[0].font.bold = True
        set_cell_background(cell_k, "F0F4F8")
        set_cell_background(cell_v, "FAFAFA")
    doc.add_paragraph()

    doc.add_heading("1. Executive Summary & Delivery Status", level=1)
    doc.add_paragraph(
        "Based on the analysis of DEV-460: Proof of Concept – SAST Integration in CI/CD Pipeline, "
        "the SAST integration has been successfully completed and validated. The primary tool chosen is Semgrep "
        "for its low scan latency (<10s) and native SARIF support. The key outcome demonstrated 100% detection of "
        "High/Critical vulnerabilities in Flask/Django test cases with <8 seconds total pipeline overhead."
    )

    doc.add_heading("2. Recommended Next Delivery Actions (Rollout Roadmap)", level=1)
    table_phases = doc.add_table(rows=4, cols=4)
    table_phases.alignment = WD_TABLE_ALIGNMENT.CENTER
    p_headers = ["Phase & Timeline", "Owner", "Key Action", "Goal"]
    for j, h in enumerate(p_headers):
        c = table_phases.cell(0, j)
        c.text = h
        c.paragraphs[0].runs[0].font.bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(c, "1F4E79")

    phases_data = [
        ("Phase 2: Multi-Service Expansion (Immediate)", "@DevOps", "Enable non-blocking SAST scans across all backend services.", "Collect baseline security telemetry without developer friction."),
        ("Phase 3: Feedback & Tuning (2-4 Weeks)", "Security Lead / @DevOps", "Review Security dashboard, tune .semgrepignore and rules.", "Eliminate false positives and prevent alert fatigue."),
        ("Phase 4: Enforcement (End of Month)", "DevOps / Eng Manager", "Transition High/Critical findings to blocking status (exit-code: 1).", "Prevent critical vulnerabilities from merging into production.")
    ]
    for row_idx, row_vals in enumerate(phases_data, start=1):
        for col_idx, val in enumerate(row_vals):
            cell = table_phases.cell(row_idx, col_idx)
            cell.text = val
            if row_idx % 2 == 0:
                set_cell_background(cell, "F2F2F2")
    doc.add_paragraph()

    doc.add_heading("3. Risk Mitigation & Blocker Assessment", level=1)
    doc.add_paragraph("• Alert Fatigue: Mitigated by maintaining a strict High/Critical focus and actionable inline remediation snippets.")
    doc.add_paragraph("• Vault Dependency: Mitigated by implementing resilient fallback handling in ci.yml and monitoring Vault cluster health.")
    doc.add_paragraph("• Scope Creep: Mitigated by sticking strictly to baseline rulesets until Phase 4 stabilizes.")
    doc.add_paragraph("• Blockers: None identified. All technical hurdles (Vault integration, scan speed, accuracy) have cleared.")

    doc.add_heading("4. Automation Architecture", level=1)
    doc.add_paragraph("• Script: tasks/sast_report_automation/send_sast_report.py (Slack, Teams, Discord dispatcher)")
    doc.add_paragraph("• Workflow: .github/workflows/sast_report_schedule.yml (Scheduled Monday 09:00 UTC)")
    doc.add_paragraph("• Vault Secrets: AppRole retrieval of webhook endpoints at secret/data/ci/sast")

    doc_path = "c:\\Internships\\devops\\tasks\\sast_report_automation\\SAST_Report_Automation_Guide.docx"
    doc.save(doc_path)
    print(f"Generated docx at: {doc_path}")

create_document()
